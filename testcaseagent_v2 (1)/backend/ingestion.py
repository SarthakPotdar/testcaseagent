"""
ingestion.py - Extracts text from PDF, XLSX, DOCX files (in-memory only).
"""

import io
import logging
from pathlib import Path

import pdfplumber
import openpyxl
from docx import Document

logger = logging.getLogger(__name__)


def extract_text_from_pdf(file_bytes: bytes) -> str:
    text_parts = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page_num, page in enumerate(pdf.pages, 1):
            text = page.extract_text()
            if text:
                text_parts.append(f"--- Page {page_num} ---\n{text}")
            for table in page.extract_tables():
                rows = [" | ".join(str(cell).strip() if cell else "" for cell in row) for row in table]
                if rows:
                    text_parts.append("\n[TABLE]\n" + "\n".join(rows) + "\n[/TABLE]")
    return "\n\n".join(text_parts)


def extract_text_from_xlsx(file_bytes: bytes) -> str:
    wb = openpyxl.load_workbook(io.BytesIO(file_bytes), data_only=True)
    text_parts = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows = []
        for row in ws.iter_rows(values_only=True):
            non_empty = [str(c).strip() for c in row if c is not None and str(c).strip()]
            if non_empty:
                rows.append(" | ".join(non_empty))
        if rows:
            text_parts.append(f"--- Sheet: {sheet_name} ---\n" + "\n".join(rows))
    return "\n\n".join(text_parts)


def extract_text_from_docx(file_bytes: bytes) -> str:
    doc = Document(io.BytesIO(file_bytes))
    text_parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text.strip())
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                text_parts.append(" | ".join(cells))
    return "\n".join(text_parts)


def ingest_file(filename: str, file_bytes: bytes) -> dict:
    ext = Path(filename).suffix.lower()
    result = {"filename": filename, "extension": ext, "content": "", "char_count": 0, "status": "success", "error": None}
    try:
        if ext == ".pdf":
            result["content"] = extract_text_from_pdf(file_bytes)
        elif ext in (".xlsx", ".xls"):
            result["content"] = extract_text_from_xlsx(file_bytes)
        elif ext == ".docx":
            result["content"] = extract_text_from_docx(file_bytes)
        else:
            result["status"] = "error"
            result["error"] = f"Unsupported file type: {ext}"
            return result

        if not result["content"].strip():
            result["status"] = "error"
            result["error"] = "No extractable text found in file."
            return result

        result["char_count"] = len(result["content"])
        logger.info(f"Ingested '{filename}': {result['char_count']} chars")
    except Exception as e:
        logger.error(f"Failed to ingest '{filename}': {e}")
        result["status"] = "error"
        result["error"] = str(e)
    return result


def ingest_multiple_files(files: list[dict]) -> list[dict]:
    """files: list of {"filename": str, "bytes": bytes}"""
    return [ingest_file(f["filename"], f["bytes"]) for f in files]
