"""
ingestion.py - Handles parsing and text extraction from uploaded files.
Supports: PDF, XLSX, DOCX
"""

import io
import logging
from pathlib import Path
from typing import Union

import pdfplumber
import openpyxl
from docx import Document

logger = logging.getLogger(__name__)


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract all text from a PDF file."""
    text_parts = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page_num, page in enumerate(pdf.pages, 1):
            text = page.extract_text()
            if text:
                text_parts.append(f"--- Page {page_num} ---\n{text}")
            # Also extract tables
            tables = page.extract_tables()
            for table in tables:
                rows = []
                for row in table:
                    cleaned = [str(cell).strip() if cell else "" for cell in row]
                    rows.append(" | ".join(cleaned))
                if rows:
                    text_parts.append("\n[TABLE]\n" + "\n".join(rows) + "\n[/TABLE]")
    return "\n\n".join(text_parts)


def extract_text_from_xlsx(file_bytes: bytes) -> str:
    """Extract all text from an XLSX file, sheet by sheet."""
    wb = openpyxl.load_workbook(io.BytesIO(file_bytes), data_only=True)
    text_parts = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows = []
        for row in ws.iter_rows(values_only=True):
            non_empty = [str(cell).strip() for cell in row if cell is not None and str(cell).strip()]
            if non_empty:
                rows.append(" | ".join(non_empty))
        if rows:
            text_parts.append(f"--- Sheet: {sheet_name} ---\n" + "\n".join(rows))
    return "\n\n".join(text_parts)


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract all text from a DOCX file."""
    doc = Document(io.BytesIO(file_bytes))
    text_parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text.strip())
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                text_parts.append(" | ".join(cells))
    return "\n".join(text_parts)


def ingest_file(filename: str, file_bytes: bytes) -> dict:
    """
    Ingest a single file and return extracted content metadata.

    Returns:
        {
            "filename": str,
            "extension": str,
            "content": str,
            "char_count": int,
            "status": "success" | "error",
            "error": str | None
        }
    """
    ext = Path(filename).suffix.lower()
    result = {
        "filename": filename,
        "extension": ext,
        "content": "",
        "char_count": 0,
        "status": "success",
        "error": None,
    }

    try:
        if ext == ".pdf":
            result["content"] = extract_text_from_pdf(file_bytes)
        elif ext in (".xlsx", ".xls"):
            result["content"] = extract_text_from_xlsx(file_bytes)
        elif ext == ".docx":
            result["content"] = extract_text_from_docx(file_bytes)
        else:
            result["status"] = "error"
            result["error"] = f"Unsupported file type: {ext}"
            return result

        if not result["content"].strip():
            result["status"] = "error"
            result["error"] = "No extractable text found in file."
            return result

        result["char_count"] = len(result["content"])
        logger.info(f"Ingested '{filename}': {result['char_count']} characters extracted.")

    except Exception as e:
        logger.error(f"Failed to ingest '{filename}': {e}")
        result["status"] = "error"
        result["error"] = str(e)

    return result


def ingest_multiple_files(files: list[dict]) -> list[dict]:
    """
    Ingest a list of files.
    Each item: {"filename": str, "bytes": bytes}
    Returns list of ingestion results.
    """
    results = []
    for f in files:
        result = ingest_file(f["filename"], f["bytes"])
        results.append(result)
    return results
